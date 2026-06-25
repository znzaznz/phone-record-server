"""Markdown 转 docx（轻量，不依赖 LibreOffice）。"""

from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT_NAME = "微软雅黑"
_FALLBACK_ENCODINGS = ("utf-8", "utf-8-sig", "gb18030", "gbk")

_SAFE_STEM = re.compile(r"^[0-9]{4}-[0-9]{2}-[0-9]{2}_[^./\\]+$")
_META_LINE = re.compile(r"^- \*\*.+\*\*:")


def validate_stem(stem: str) -> str:
    stem = (stem or "").strip()
    if not stem or ".." in stem or "/" in stem or "\\" in stem:
        raise ValueError("invalid stem")
    if not _SAFE_STEM.match(stem):
        raise ValueError("invalid stem format")
    return stem


def _set_rfonts(rPr, name: str) -> None:
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), name)


def _set_lang(rPr) -> None:
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-CN")


def _configure_styles(doc: Document) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = FONT_NAME
        rPr = style.element.get_or_add_rPr()
        _set_rfonts(rPr, FONT_NAME)
        _set_lang(rPr)


def _font_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        rPr = run._element.get_or_add_rPr()
        _set_rfonts(rPr, FONT_NAME)
        _set_lang(rPr)


def _read_md_text(md_path: Path) -> str:
    data = md_path.read_bytes()
    for enc in _FALLBACK_ENCODINGS:
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _content_lines(md_text: str) -> list[str]:
    """去掉 task_id/model 等元数据，只保留标题与正文。"""
    lines = md_text.splitlines()
    title: str | None = None
    body: list[str] = []
    past_sep = False
    for raw in lines:
        line = raw.rstrip()
        if not past_sep:
            if line.startswith("# "):
                title = line[2:].strip()
                if title.endswith(" — 总结"):
                    title = title[:-5].strip()
                continue
            if line.strip() == "---":
                past_sep = True
                continue
            if _META_LINE.match(line.strip()):
                continue
            if not line.strip():
                continue
            past_sep = True
            body.append(line)
            continue
        body.append(line)
    out: list[str] = []
    if title:
        out.append(f"# {title}")
    out.extend(body)
    return out


def md_to_docx_bytes(md_text: str) -> bytes:
    doc = Document()
    _configure_styles(doc)
    for line in _content_lines(md_text):
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("---"):
            continue
        else:
            p = doc.add_paragraph(line)
        _font_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def md_file_to_docx_bytes(md_path: Path) -> bytes:
    return md_to_docx_bytes(_read_md_text(md_path))